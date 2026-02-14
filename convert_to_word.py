#!/usr/bin/env python3
"""
Convert PROJECT_PLAN.md to a Word document
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set hyperlink style
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def parse_markdown_to_doc(md_file, doc_file):
    """Convert markdown file to Word document"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines (but keep some spacing)
        if not line:
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            i += 1
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            i += 1
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            i += 1
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            i += 1
        elif line.startswith('##### '):
            p = doc.add_heading(line[6:], level=5)
            i += 1
        elif line.startswith('###### '):
            p = doc.add_heading(line[7:], level=6)
            i += 1
        
        # Horizontal rule
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('_' * 80).font.size = Pt(1)
            i += 1
        
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run(''.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Check for checkbox
            if text.startswith('[ ]'):
                text = '☐ ' + text[3:].strip()
            elif text.startswith('[x]') or text.startswith('[X]'):
                text = '☑ ' + text[3:].strip()
            
            p = doc.add_paragraph(text, style='List Bullet')
            # Handle bold text
            p = format_text(p, text)
            i += 1
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
            p = format_text(p, text)
            i += 1
        
        # Tables (basic detection)
        elif '|' in line and '---' in lines[i+1] if i+1 < len(lines) else False:
            # Parse table
            headers = [h.strip() for h in line.split('|')[1:-1]]
            i += 2  # Skip header and separator
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Grid Accent 1'
            
            # Add headers
            hdr_cells = table.rows[0].cells
            for j, header in enumerate(headers):
                hdr_cells[j].text = header
                hdr_cells[j].paragraphs[0].runs[0].bold = True
            
            # Add rows
            while i < len(lines) and '|' in lines[i]:
                row_data = [d.strip() for d in lines[i].split('|')[1:-1]]
                row_cells = table.add_row().cells
                for j, data in enumerate(row_data):
                    if j < len(row_cells):
                        row_cells[j].text = data
                i += 1
        
        # Regular paragraphs
        else:
            p = doc.add_paragraph()
            p = format_text(p, line)
            i += 1
    
    doc.save(doc_file)
    print(f"Word document created: {doc_file}")

def format_text(paragraph, text):
    """Format text with bold, italic, and links"""
    # Clear existing text
    paragraph.clear()
    
    # Pattern for bold **text**
    bold_pattern = r'\*\*(.*?)\*\*'
    # Pattern for italic *text*
    italic_pattern = r'(?<!\*)\*([^*]+?)\*(?!\*)'
    # Pattern for links [text](url)
    link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
    
    # Find all formatting
    positions = []
    
    # Find bold
    for match in re.finditer(bold_pattern, text):
        positions.append((match.start(), match.end(), 'bold', match.group(1)))
    
    # Find italic (but not bold)
    for match in re.finditer(italic_pattern, text):
        # Check if it's not part of bold
        is_bold = False
        for pos in positions:
            if pos[3] == 'bold' and match.start() >= pos[0] and match.end() <= pos[1]:
                is_bold = True
                break
        if not is_bold:
            positions.append((match.start(), match.end(), 'italic', match.group(1)))
    
    # Find links
    for match in re.finditer(link_pattern, text):
        positions.append((match.start(), match.end(), 'link', (match.group(1), match.group(2))))
    
    # Sort by position
    positions.sort(key=lambda x: x[0])
    
    # Build formatted text
    last_pos = 0
    for start, end, fmt_type, content in positions:
        # Add text before this format
        if start > last_pos:
            paragraph.add_run(text[last_pos:start])
        
        # Add formatted text
        if fmt_type == 'bold':
            run = paragraph.add_run(content)
            run.bold = True
        elif fmt_type == 'italic':
            run = paragraph.add_run(content)
            run.italic = True
        elif fmt_type == 'link':
            # For links, just add the text (Word hyperlinks are complex)
            run = paragraph.add_run(content[0])
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True
        
        last_pos = end
    
    # Add remaining text
    if last_pos < len(text):
        paragraph.add_run(text[last_pos:])
    
    return paragraph

if __name__ == '__main__':
    import sys
    md_file = 'PROJECT_PLAN.md'
    doc_file = 'PROJECT_PLAN.docx'
    
    try:
        parse_markdown_to_doc(md_file, doc_file)
    except ImportError:
        print("Error: python-docx library not found.")
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
        print("Retrying...")
        parse_markdown_to_doc(md_file, doc_file)
